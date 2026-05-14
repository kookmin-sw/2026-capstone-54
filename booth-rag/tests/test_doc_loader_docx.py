from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from booth_rag.ingestion.doc_loader import chunk_document, load_document, parse_docx_sections


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_heading("개요", level=1)
    doc.add_paragraph("미핏은 캡스톤 54팀의 면접 준비 플랫폼입니다.")
    doc.add_paragraph("이력서 기반 AI 면접을 제공합니다.")
    doc.add_heading("시스템 구성", level=1)
    doc.add_heading("백엔드", level=2)
    doc.add_paragraph("Django 6.0, DRF, Celery, Channels 를 사용합니다.")
    doc.add_paragraph("LangChain + OpenAI 로 면접 질문을 생성합니다.")
    doc.add_heading("프론트엔드", level=2)
    doc.add_paragraph("React 19, Vite, Three.js, GSAP 기반.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "항목"
    table.rows[0].cells[1].text = "값"
    table.rows[1].cells[0].text = "면접 횟수"
    table.rows[1].cells[1].text = "10K+"

    path = tmp_path / "report.docx"
    doc.save(str(path))
    return path


def test_parse_docx_sections_splits_by_headings(sample_docx: Path):
    sections = parse_docx_sections(sample_docx)
    assert len(sections) >= 3
    titles = [s.heading_path[-1] if s.heading_path else "" for s in sections]
    assert "개요" in titles
    assert "백엔드" in titles
    assert "프론트엔드" in titles


def test_parse_docx_sections_track_heading_path(sample_docx: Path):
    sections = parse_docx_sections(sample_docx)
    backend = next(s for s in sections if "백엔드" in s.heading_path)
    assert "시스템 구성" in backend.heading_path
    assert "Django" in backend.text


def test_chunk_document_emits_docx_sections(sample_docx: Path):
    doc = load_document(sample_docx)
    assert doc is not None
    assert doc.docx_sections, "docx_sections should be populated"
    chunks = chunk_document(doc)
    assert all(c.kind == "docx_section" for c in chunks)
    assert any("개요" in c.symbol for c in chunks)
    assert any("백엔드" in c.symbol for c in chunks)
    backend_chunk = next(c for c in chunks if "백엔드" in (c.symbol or ""))
    assert "Django" in backend_chunk.text
    assert backend_chunk.text.startswith("# ")
