from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument

from booth_rag.ingestion.code_chunker import CodeChunk
from booth_rag.utils.secret_filter import is_secret_file, redact_secrets

logger = logging.getLogger(__name__)

_TARGET_CHARS = 1500
_SOFT_CAP_CHARS = 1800
_HARD_CAP_CHARS = 2400
_PARA_BUFFER_FLUSH = 200
_HEADING_PATTERN = re.compile(r"^Heading\s*(\d+)", re.IGNORECASE)


@dataclass(frozen=True)
class DocxSection:
    heading_path: tuple[str, ...]
    text: str
    paragraph_count: int


@dataclass(frozen=True)
class LoadedDoc:
    rel_path: str
    text: str
    docx_sections: tuple[DocxSection, ...] = ()


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    m = _HEADING_PATTERN.match(style_name.strip())
    if not m:
        return None
    try:
        level = int(m.group(1))
    except ValueError:
        return None
    return level if 1 <= level <= 6 else None


def _iter_docx_blocks(doc) -> list[tuple[str, str | None, str]]:
    blocks: list[tuple[str, str | None, str]] = []
    for para in doc.paragraphs:
        text = para.text.strip() if para.text else ""
        if not text:
            continue
        style = para.style.name if para.style else None
        blocks.append(("paragraph", style, text))

    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(("table", None, "\n".join(rows)))

    return blocks


def _flush_section(
    out: list[DocxSection],
    heading_stack: list[str],
    buffer: list[str],
    para_count: int,
) -> None:
    if not buffer:
        return
    text = "\n\n".join(buffer).strip()
    if not text:
        return
    out.append(
        DocxSection(
            heading_path=tuple(heading_stack),
            text=text,
            paragraph_count=para_count,
        )
    )


def parse_docx_sections(path: Path) -> list[DocxSection]:
    """Parse a docx into heading-aware sections without loading entire text in memory twice."""
    doc = DocxDocument(str(path))
    sections: list[DocxSection] = []
    heading_stack: list[str] = []
    buffer: list[str] = []
    para_count = 0

    for kind, style, content in _iter_docx_blocks(doc):
        if kind == "paragraph":
            level = _heading_level(style)
            if level is not None:
                _flush_section(sections, heading_stack, buffer, para_count)
                buffer = []
                para_count = 0
                heading_stack = [*heading_stack[: level - 1], content]
                continue
            buffer.append(content)
            para_count += 1
        else:
            buffer.append(content)
            para_count += 1

        joined_len = sum(len(b) + 2 for b in buffer)
        if joined_len > _HARD_CAP_CHARS:
            _flush_section(sections, heading_stack, buffer, para_count)
            buffer = []
            para_count = 0

    _flush_section(sections, heading_stack, buffer, para_count)
    return sections


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _split_long(text: str) -> list[str]:
    if len(text) <= _SOFT_CAP_CHARS:
        return [text]
    parts: list[str] = []
    pos = 0
    while pos < len(text):
        end = min(pos + _TARGET_CHARS, len(text))
        if end < len(text):
            window_end = min(end + _PARA_BUFFER_FLUSH, len(text))
            br = text.rfind("\n", end - _PARA_BUFFER_FLUSH, window_end)
            if br > pos:
                end = br
        parts.append(text[pos:end].strip())
        pos = end
    return [p for p in parts if p]


def load_document(path: Path, rel_root: Path | None = None) -> LoadedDoc | None:
    if not path.is_file():
        return None
    if is_secret_file(path.name, str(path)):
        return None
    suf = path.suffix.lower()
    try:
        if suf in {".md", ".mdx", ".markdown", ".txt", ".rst"}:
            text = _read_text(path)
            rel = str(path.relative_to(rel_root)) if rel_root else path.name
            return LoadedDoc(rel_path=rel, text=text)
        if suf == ".docx":
            sections = parse_docx_sections(path)
            text = "\n\n".join(s.text for s in sections)
            rel = str(path.relative_to(rel_root)) if rel_root else path.name
            return LoadedDoc(rel_path=rel, text=text, docx_sections=tuple(sections))
    except Exception as exc:
        logger.warning("Failed to load document %s: %s", path, exc)
        return None
    return None


def _chunks_from_text(rel_path: str, text: str, suffix: str) -> list[CodeChunk]:
    from booth_rag.ingestion.code_chunker import chunk_file
    from booth_rag.ingestion.code_walker import CodeFile

    pseudo = CodeFile(
        path=Path(rel_path),
        rel_path=rel_path,
        suffix=suffix,
        size_bytes=len(text.encode("utf-8")),
        text=text,
    )
    return chunk_file(pseudo)


def chunk_document(doc: LoadedDoc) -> list[CodeChunk]:
    if doc.docx_sections:
        chunks: list[CodeChunk] = []
        idx = 0
        for section in doc.docx_sections:
            heading_path = " > ".join(section.heading_path) if section.heading_path else "(no heading)"
            redacted = redact_secrets(section.text).text
            for part in _split_long(redacted):
                if not part.strip():
                    continue
                body = f"# {heading_path}\n\n{part}"
                chunks.append(
                    CodeChunk(
                        rel_path=doc.rel_path,
                        chunk_index=idx,
                        text=body,
                        kind="docx_section",
                        symbol=heading_path[:120],
                        line_start=0,
                        line_end=section.paragraph_count,
                    )
                )
                idx += 1
        return chunks

    suffix = ".md" if doc.rel_path.lower().endswith((".md", ".mdx", ".markdown")) else ".txt"
    return _chunks_from_text(doc.rel_path, doc.text, suffix)
